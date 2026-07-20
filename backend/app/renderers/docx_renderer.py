"""The editable DOCX renderer (Renderer Track R1, Phases 3–6).

``DocxRenderer`` is the first concrete :class:`~app.renderers.protocol.Renderer`.
It turns a persisted :class:`~app.renderers.snapshot.ReportSnapshot` into a valid,
editable ``.docx`` document using ``python-docx`` and nothing else.

It obeys the rendering invariant in full (`PLATFORM_ARCHITECTURE.md` §2.2): no
LLM, no retrieval, no evidence scoring, no domain reasoning, no live document
reads. It only reads the snapshot it is given and writes bytes.

Editability: every heading is a real Word heading, every register is a real Word
table, formatting flows from named styles, and text is real text — never a
screenshot. There are no macros and no external relationships.

Determinism (Phase 6): the visible content and structure are a pure function of
the snapshot and options. python-docx writes a ZIP whose entry timestamps come
from the wall clock, so the produced container is re-packed with pinned
timestamps and normalized entry attributes; within one python/zlib build this
yields byte-identical output for identical inputs. A container-independent
``semantic_digest`` over the canonical input plus the normalized document body is
also emitted, so logical identity is provable even where the container bytes
could wobble across library builds.
"""

from __future__ import annotations

import dataclasses
import hashlib
import io
import json
import zipfile
from datetime import datetime, timezone
from typing import Any, Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.renderers import docx_styles as S
from app.renderers.protocol import RenderedArtifact, RendererError, RendererOptions
from app.renderers.sanitize import clean_text, safe_filename
from app.renderers.snapshot import ReportSnapshot

CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
RENDERER_ID = "docx-renderer"
RENDERER_VERSION = "docx-renderer-v1"

# A pinned, non-wall-clock fallback for the document's core-property dates when
# the persisted report carries no parseable ``generatedAt``. Constant, so the
# output stays deterministic; never ``datetime.now()``.
PINNED_EPOCH = datetime(2001, 1, 1, 0, 0, 0, tzinfo=timezone.utc)

# ZIP DOS epoch — the smallest timestamp a ZIP entry can carry — pinned onto
# every entry so the container has no wall-clock time anywhere.
_ZIP_PINNED_TIME = (1980, 1, 1, 0, 0, 0)


class DocxRenderer:
    """Pure snapshot → editable DOCX transformation."""

    renderer_id = RENDERER_ID
    renderer_version = RENDERER_VERSION
    content_type = CONTENT_TYPE

    def render(self, snapshot: ReportSnapshot, options: RendererOptions) -> RenderedArtifact:
        options = options.normalized()
        document = Document()
        S.apply_styles(document)

        section = document.sections[0]
        self._page_setup(section, options)
        self._core_properties(document, snapshot)
        self._header_footer(document, section, snapshot)

        self._cover(document, snapshot)
        if options.include_table_of_contents:
            self._table_of_contents(document)
        self._executive_summary(document, snapshot)
        self._analysis_overview(document, snapshot)
        if snapshot.report.workflow_steps:
            self._workflow(document, snapshot)
        if snapshot.report.risks:
            self._risk_register(document, snapshot)
        if snapshot.report.suggested_actions:
            self._recommendations(document, snapshot)
        self._citations(document, snapshot, options)
        if options.include_audit_appendix:
            self._audit_appendix(document, snapshot)

        raw = io.BytesIO()
        document.save(raw)
        final_bytes = _normalize_container(raw.getvalue())

        if len(final_bytes) > options.max_output_bytes:
            raise RendererError(
                "export_too_large",
                "The rendered document exceeded the maximum allowed size.",
            )

        content_hash = hashlib.sha256(final_bytes).hexdigest()
        semantic_digest = _semantic_digest(snapshot, options, final_bytes)
        filename = safe_filename(
            snapshot.report.custom_persona or snapshot.report.persona,
            snapshot.report.market,
            snapshot.report.id,
        )

        return RenderedArtifact(
            data=final_bytes,
            filename=filename,
            content_type=self.content_type,
            renderer_id=self.renderer_id,
            renderer_version=self.renderer_version,
            content_hash=content_hash,
            semantic_digest=semantic_digest,
            byte_size=len(final_bytes),
            telemetry={
                "workflowSteps": len(snapshot.report.workflow_steps),
                "risks": len(snapshot.report.risks),
                "citations": len(snapshot.report.citations),
                "sourceVersions": len(snapshot.audit.source_versions),
                "evidenceBindings": len(snapshot.audit.evidence_bindings),
                "auditPresent": snapshot.audit.present,
            },
        )

    # --- page + document chrome -------------------------------------------

    def _page_setup(self, section, options: RendererOptions) -> None:
        if options.page_size == "Letter":
            section.page_width = Mm(215.9)
            section.page_height = Mm(279.4)
        else:  # A4
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        for attr in ("top_margin", "bottom_margin"):
            setattr(section, attr, Mm(22))
        for attr in ("left_margin", "right_margin"):
            setattr(section, attr, Mm(22))
        section.header_distance = Mm(12)
        section.footer_distance = Mm(12)

    def _core_properties(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        core = document.core_properties
        pinned = report.generated_at or PINNED_EPOCH
        # python-docx serializes these as naive UTC; strip tzinfo deterministically.
        naive = pinned.astimezone(timezone.utc).replace(tzinfo=None)
        persona = report.custom_persona or report.persona or "Report"
        core.title = clean_text(f"Evidentia report — {persona}", limit=240)
        core.author = "Evidentia"
        core.last_modified_by = f"Evidentia {RENDERER_VERSION}"
        core.subject = clean_text(snapshot.tenant.company_name or "Evidentia report", limit=240)
        core.category = clean_text(report.category, limit=120)
        core.comments = f"Rendered by {RENDERER_VERSION}. Deterministic projection of a persisted Evidentia report."
        core.keywords = "Evidentia; persona report; evidence-grounded"
        core.created = naive
        core.modified = naive
        core.revision = 1

    def _header_footer(self, document, section, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.style = document.styles[S.S_METADATA]
        persona = report.custom_persona or report.persona or "Report"
        _run(hp, f"EVIDENTIA · {snapshot.tenant.company_name or report.company} · {persona}")

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.style = document.styles[S.S_METADATA]
        content_width = section.page_width - section.left_margin - section.right_margin
        fp.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
        _run(fp, "Evidentia · Confidential")
        fp.add_run().add_tab()
        _run(fp, "Page ")
        _add_field(fp, " PAGE ")
        _run(fp, " of ")
        _add_field(fp, " NUMPAGES ")

    # --- content sections -------------------------------------------------

    def _cover(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        _para(document, "EVIDENTIA", S.S_COVER_LABEL)
        _para(document, "Evidence-grounded persona report", "Subtitle")
        _spacer(document, 10)
        persona = report.custom_persona or report.persona or "Persona report"
        _para(document, persona, "Title")
        subtitle_bits = [b for b in (snapshot.tenant.company_name or report.company, report.market) if b]
        if subtitle_bits:
            _para(document, " · ".join(subtitle_bits), "Subtitle")
        _spacer(document, 16)

        # A compact metadata block on the cover. Only persisted facts appear;
        # absent facts are omitted, never guessed.
        rows: List[tuple[str, str]] = []
        if report.market:
            rows.append(("Market", report.market))
        base_persona = report.persona
        if base_persona:
            rows.append(("Persona", base_persona))
        if report.custom_persona:
            rows.append(("Custom role", report.custom_persona))
        if report.category:
            rows.append(("Category", report.category))
        rows.append(("Generated", _format_stamp(report)))
        rows.append(("Corpus", _corpus_label(snapshot)))
        rows.append(("Generation", _generation_label(report)))
        if report.confidence and not report.analytical_output_empty:
            rows.append(("Baseline score", f"{report.confidence}% document-count heuristic"))
        elif report.analytical_output_empty:
            rows.append(("Analytical confidence", "N/A — no accepted claims to score"))
        rows.append(("Renderer", RENDERER_VERSION))
        rows.append(("Report id", report.id or "—"))
        _definition_table(document, rows)

        _page_break(document)

    def _table_of_contents(self, document) -> None:
        _para(document, "Contents", "Heading 1")
        p = document.add_paragraph(style=document.styles["Normal"])
        _toc_field(p)
        _page_break(document)

    def _executive_summary(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        _para(document, "Executive summary", "Heading 1")
        if report.summary:
            _para(document, report.summary, S.S_BODY)
        else:
            _para(document, "No executive summary was recorded for this report.", S.S_BODY)

        if report.top_finding:
            _para(document, "Top finding", "Heading 2")
            _para(document, report.top_finding, S.S_BODY)

        metrics = report.metrics
        if metrics.present:
            _para(document, "Headline metrics", "Heading 2")
            if report.analytical_output_empty:
                rows = [
                    ("Frozen source versions", str(snapshot.audit.source_version_count) if snapshot.audit.present else "—"),
                    ("Selected evidence sections", str(snapshot.audit.evidence_section_count) if snapshot.audit.present else "—"),
                    ("Source bindings", str(len(snapshot.audit.evidence_bindings))),
                    ("Analytical confidence", "N/A — no accepted claims to score"),
                ]
            else:
                rows = [
                    ("Documents analyzed", str(metrics.documents_analyzed)),
                    ("Citations used", str(metrics.citations_used)),
                    ("Risks flagged", str(metrics.risks_flagged)),
                    ("Workflow steps", str(len(report.workflow_steps))),
                    ("Baseline score", f"{metrics.confidence}% document-count heuristic"),
                ]
            _definition_table(document, rows)

    def _analysis_overview(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        audit = snapshot.audit
        _para(document, "Analysis overview", "Heading 1")

        brief = report.persona_brief
        if brief.description:
            _para(document, "Configured persona context" if report.analytical_output_empty else "Persona brief", "Heading 2")
            if report.analytical_output_empty:
                _para(
                    document,
                    "User-selected configuration used to scope the analysis; this is not an evidence-derived finding.",
                    S.S_BODY,
                )
            _para(document, brief.description, S.S_BODY)
        if brief.priorities:
            _para(document, "Priorities", "Heading 3")
            for item in brief.priorities:
                _bullet(document, item)
        if brief.goals:
            _para(document, "Goals", "Heading 3")
            for item in brief.goals:
                _bullet(document, item)
        if brief.output_style:
            _para(document, "Output style", "Heading 3")
            _para(document, brief.output_style, S.S_BODY)

        _para(document, "Scope & methodology", "Heading 2")
        scope_rows = [
            ("Corpus mode", _corpus_label(snapshot)),
            ("Documents analyzed", str(report.metrics.documents_analyzed)),
            ("Frozen source versions", str(audit.source_version_count) if audit.present else "—"),
            ("Selected evidence sections", str(audit.evidence_section_count) if audit.present else "—"),
            ("Generation mode", _generation_label(report)),
            ("Retrieval engine", audit.retrieval_engine_version or "—"),
        ]
        _definition_table(document, scope_rows)
        _para(
            document,
            "Evidentia's deterministic pipeline produces a complete, evidence-grounded "
            "report with no language model. When enabled, an LLM only refines that "
            "baseline and every claim remains bound to a cited source section. This "
            "document is a deterministic rendering of the persisted report and its "
            "report-local source audit; it performs no retrieval or reasoning of its own.",
            S.S_BODY,
        )

    def _workflow(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        _para(document, "Recommended workflow", "Heading 1")
        if not report.workflow_steps:
            _para(
                document,
                "No workflow steps could be grounded in the selected documents.",
                S.S_BODY,
            )
            return
        for step in report.workflow_steps:
            heading = f"Step {step.step}"
            if step.title:
                heading = f"{heading} · {step.title}"
            _para(document, heading, "Heading 2")
            if step.description:
                _para(document, step.description, S.S_BODY)
            if step.why_it_matters:
                _labelled(document, "Why it matters", step.why_it_matters)
            if step.expected_output:
                _labelled(document, "Expected output", step.expected_output)
            _evidence_line(document, snapshot, step.evidence_code)

    def _risk_register(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        _para(document, "Risk register", "Heading 1")
        if not report.risks:
            _para(
                document,
                "No risks met the evidence-support threshold for this corpus.",
                S.S_BODY,
            )
            return

        headers = ["Severity", "Risk", "Business impact", "Mitigation", "Owner", "Evidence"]
        table = _new_table(document, headers)
        for risk in report.risks:
            cells = table.add_row().cells
            _cell(cells[0], risk.severity or "—", style=S.severity_style(risk.severity))
            if risk.severity:
                _shade_cell(cells[0], S.severity_fill(risk.severity))
            _cell(cells[1], risk.title or "—", bold=True)
            _cell(cells[2], risk.business_impact or "—")
            _cell(cells[3], risk.recommended_fix or "—")
            _cell(cells[4], _owner_label(risk))
            _cell(cells[5], _evidence_label(risk.evidence_code), style=S.S_CITATION)

    def _recommendations(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        _para(document, "Recommendations & next actions", "Heading 1")
        if not report.suggested_actions:
            _para(document, "No suggested actions were recorded for this report.", S.S_BODY)
            return
        headers = ["#", "Action", "Detail"]
        table = _new_table(document, headers, widths_mm=(10, 55, 100))
        for index, action in enumerate(report.suggested_actions, start=1):
            cells = table.add_row().cells
            _cell(cells[0], str(index))
            _cell(cells[1], action.title or "—", bold=True)
            _cell(cells[2], action.detail or "—")

    def _citations(self, document, snapshot: ReportSnapshot, options: RendererOptions) -> None:
        report = snapshot.report
        audit = snapshot.audit
        _para(document, "Evidence & citations", "Heading 1")

        if not report.citations:
            _para(document, "No source citations were bound for this report.", S.S_BODY)
        else:
            _para(
                document,
                "Citations bound to this report's frozen source audit show their "
                "document, version, section and evidence exactly as captured in the "
                "report-local M4 snapshot — never from any current document version.",
                S.S_BODY,
            )
            # Legacy demo rows may carry only the persisted report record with no
            # source-audit binding. We preserve their excerpt as honest
            # compatibility ONLY when the corpus mode is *explicitly* demo; for a
            # tenant (or unknown) corpus, an unbound excerpt is never presented as
            # frozen evidence.
            demo_corpus = audit.corpus_mode == "demo"
            for citation in report.citations:
                binding = audit.binding_for(citation.id)
                if binding is not None:
                    # Fully bound: citation id, document title, version, section and
                    # the evidence quote ALL come from the same frozen audit record.
                    # report_json's excerpt/source are never mixed in here.
                    _para(document, binding.citation_id or citation.id or "Citation", "Heading 3")
                    title = binding.document_title or binding.original_filename
                    if title:
                        _para(document, title, S.S_BODY)
                    meta_bits: List[str] = []
                    if binding.document_version_id:
                        meta_bits.append(f"Version {binding.document_version_id}")
                    meta_bits.append(f"Section {binding.section_ordinal + 1}")
                    if binding.heading_path:
                        meta_bits.append(" › ".join(binding.heading_path))
                    meta_bits.append("cited in final" if binding.cited_in_final else "selected as evidence")
                    _para(document, " · ".join(meta_bits), S.S_CITATION)
                    if options.include_evidence_excerpts and binding.excerpt:
                        _para(document, f"“{binding.excerpt}”", S.S_EVIDENCE_QUOTE)
                elif demo_corpus:
                    # Explicit demo corpus, no binding: the persisted report record
                    # is the only source. Show it, but label it honestly — never as
                    # frozen source-audit evidence.
                    _para(document, citation.id or "Citation", "Heading 3")
                    if citation.source:
                        _para(document, citation.source, S.S_BODY)
                    if citation.section:
                        _para(document, citation.section, S.S_CITATION)
                    _para(
                        document,
                        "From the persisted report record — no frozen source-audit "
                        "binding exists for this demo citation.",
                        S.S_CITATION,
                    )
                    if options.include_evidence_excerpts and citation.excerpt:
                        _para(document, f"“{citation.excerpt}”", S.S_EVIDENCE_QUOTE)
                else:
                    # Tenant (or unknown) corpus, no binding: do NOT present the
                    # report_json excerpt as frozen evidence. Say the source audit is
                    # unavailable and omit the evidence quote entirely.
                    _para(document, citation.id or "Citation", "Heading 3")
                    if citation.source:
                        _para(document, citation.source, S.S_BODY)
                    _para(
                        document,
                        "Source audit unavailable for this citation; evidence quote omitted.",
                        S.S_CITATION,
                    )
                if citation.why_it_matters:
                    _labelled(document, "Why it matters", citation.why_it_matters)

        # A tabular source appendix straight from the M4 audit — every frozen
        # version and every selected evidence section, exactly as persisted.
        if audit.present and audit.evidence_bindings:
            _para(document, "Source appendix", "Heading 2")
            headers = ["Citation", "Document", "Version", "Section", "Rank", "Cited"]
            table = _new_table(document, headers, widths_mm=(24, 46, 30, 40, 12, 14))
            for binding in audit.evidence_bindings:
                cells = table.add_row().cells
                _cell(cells[0], binding.citation_id or "—", style=S.S_CITATION)
                _cell(cells[1], binding.document_title or binding.original_filename or "—")
                _cell(cells[2], binding.document_version_id or "—", style=S.S_CITATION)
                section_label = binding.section_title or "—"
                if binding.heading_path:
                    section_label = f"{section_label}\n{' › '.join(binding.heading_path)}"
                _cell(cells[3], section_label)
                _cell(cells[4], str(binding.retrieval_rank))
                _cell(cells[5], "Yes" if binding.cited_in_final else "No")

    def _audit_appendix(self, document, snapshot: ReportSnapshot) -> None:
        report = snapshot.report
        audit = snapshot.audit
        _page_break(document)
        _para(document, "Audit appendix", "Heading 1")
        if not audit.present:
            _para(
                document,
                "No report-local source audit was available for this report. It was "
                "most likely produced before source-audit provenance was recorded; the "
                "report content above is shown exactly as persisted.",
                S.S_BODY,
            )
        rows = [
            ("Corpus mode", _corpus_label(snapshot)),
            ("Snapshot digest", audit.corpus_snapshot_digest or "—"),
            ("Retrieval engine", audit.retrieval_engine_version or "—"),
            ("Orchestrator", audit.orchestrator_version or "—"),
            ("Execution mode", audit.execution_mode or report.generation_mode or "—"),
            ("LLM provider", audit.llm_provider or report.llm_provider or "none"),
            ("LLM model", audit.llm_model or report.llm_model or "—"),
            ("Frozen source versions", str(audit.source_version_count)),
            ("Selected evidence sections", str(audit.evidence_section_count)),
            ("Generation status", audit.generation_status or "—"),
            ("Renderer", f"{RENDERER_ID} · {RENDERER_VERSION}"),
        ]
        _definition_table(document, rows)

        if audit.source_versions:
            _para(document, "Source versions", "Heading 2")
            headers = ["#", "Document", "Version", "No.", "Manifest", "Finalization"]
            table = _new_table(document, headers, widths_mm=(8, 38, 34, 10, 38, 38))
            for version in audit.source_versions:
                cells = table.add_row().cells
                _cell(cells[0], str(version.position + 1))
                _cell(cells[1], version.document_id or "—", style=S.S_CITATION)
                _cell(cells[2], version.document_version_id or "—", style=S.S_CITATION)
                _cell(cells[3], str(version.version_no))
                _cell(cells[4], _short_hash(version.manifest_sha256), style=S.S_CITATION)
                _cell(cells[5], _short_hash(version.finalization_target_digest), style=S.S_CITATION)


# --------------------------------------------------------------------------- #
# Low-level document helpers (module-level; no renderer state)
# --------------------------------------------------------------------------- #


def _run(paragraph, text: str, *, bold: Optional[bool] = None, italic: Optional[bool] = None):
    run = paragraph.add_run(clean_text(text))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _para(container, text: str, style: str):
    return container.add_paragraph(clean_text(text), style=style)


def _bullet(document, text: str):
    return document.add_paragraph(clean_text(text), style="List Bullet")


def _labelled(document, label: str, text: str) -> None:
    p = document.add_paragraph(style=S.S_BODY)
    lead = p.add_run(clean_text(f"{label}: "))
    lead.bold = True
    p.add_run(clean_text(text))


def _spacer(document, points: float) -> None:
    p = document.add_paragraph(style=document.styles["Normal"])
    p.paragraph_format.space_after = Pt(points)


def _page_break(document) -> None:
    document.add_page_break()


def _owner_label(risk) -> str:
    if risk.is_insufficient:
        return "documentation gap"
    return risk.owner or "—"


def _evidence_label(code: str) -> str:
    text = (code or "").strip()
    if not text:
        return "—"
    if text.upper() == "N/A":
        return "Insufficient evidence"
    return text


def _evidence_line(document, snapshot: ReportSnapshot, code: str) -> None:
    binding = snapshot.audit.binding_for(code) if code else None
    label = _evidence_label(code)
    if binding is not None:
        section = f"Section {binding.section_ordinal + 1}"
        label = f"{code} · {binding.document_title or '—'} · {section}"
    p = document.add_paragraph(style=S.S_CITATION)
    lead = p.add_run("Evidence: ")
    lead.bold = True
    p.add_run(clean_text(label))


def _definition_table(document, rows: Iterable[tuple[str, str]]) -> None:
    materialized = [(k, v) for k, v in rows]
    if not materialized:
        return
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for key, value in materialized:
        cells = table.add_row().cells
        _cell(cells[0], key, bold=True, style=S.S_TABLE_CELL)
        _shade_cell(cells[0], "F2F2F4")
        _cell(cells[1], value, style=S.S_TABLE_CELL)
    _space_after_table(document)


def _new_table(document, headers: List[str], *, widths_mm: Optional[tuple[float, ...]] = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = widths_mm is None
    header_row = table.rows[0]
    _repeat_header(header_row)
    for index, label in enumerate(headers):
        cell = header_row.cells[index]
        _cell(cell, label, style=S.S_TABLE_HEADER)
        _shade_cell(cell, "1A1A1C")
    if widths_mm is not None:
        for index, width in enumerate(widths_mm):
            for row in table.rows:
                row.cells[index].width = Mm(width)
    return table


def _cell(cell, text: str, *, bold: Optional[bool] = None, style: str = S.S_TABLE_CELL) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    # python-docx resolves a paragraph style by name; the style is registered in
    # docx_styles.apply_styles, so this never falls back to Normal silently.
    paragraph.style = style
    # Multi-line cell content becomes multiple runs separated by line breaks.
    segments = clean_text(text).split("\n")
    for index, segment in enumerate(segments):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(segment)
        if bold is not None:
            run.bold = bold


def _space_after_table(document) -> None:
    document.add_paragraph(style=document.styles["Normal"]).paragraph_format.space_after = Pt(4)


def _short_hash(value: str) -> str:
    text = (value or "").strip()
    if not text:
        return "—"
    return text if len(text) <= 20 else f"{text[:16]}…"


def _corpus_label(snapshot: ReportSnapshot) -> str:
    if not snapshot.audit.present:
        return "Unavailable"
    return {"tenant": "Tenant corpus", "demo": "Sample corpus"}.get(
        snapshot.audit.corpus_mode, snapshot.audit.corpus_mode or "Unavailable"
    )


def _generation_label(report) -> str:
    label = report.mode_label
    if report.is_llm and report.llm_model:
        return f"{label} ({report.llm_provider or 'llm'} · {report.llm_model})"
    return label


def _format_stamp(report) -> str:
    if report.generated_at is not None:
        dt = report.generated_at
        return dt.strftime("%d %b %Y · %H:%M UTC")
    return report.generated_at_raw or "—"


# --- OXML field / table helpers -------------------------------------------


def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field (e.g. ``PAGE``/``NUMPAGES``) into ``paragraph``."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run = paragraph.add_run()
        run._r.append(element)


def _toc_field(paragraph) -> None:
    """A real TOC field. Word offers to update it; until then a placeholder line
    stands in (honest, editable — never a screenshot)."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run1 = paragraph.add_run()
    run1._r.append(begin)
    run2 = paragraph.add_run()
    run2._r.append(instr)
    run3 = paragraph.add_run()
    run3._r.append(separate)
    paragraph.add_run('Right-click and choose "Update Field" to build the table of contents.')
    run4 = paragraph.add_run()
    run4._r.append(end)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


# --- determinism -----------------------------------------------------------


def _normalize_container(data: bytes) -> bytes:
    """Re-pack the DOCX ZIP with pinned timestamps and normalized attributes.

    python-docx builds entry order deterministically for identical documents, so
    order is preserved (keeping ``[Content_Types].xml`` first for strict
    readers). Only the wall-clock timestamps, host-OS attributes and extra fields
    — the sources of container non-determinism — are normalized away.
    """
    source = zipfile.ZipFile(io.BytesIO(data), "r")
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as dest:
        for name in source.namelist():
            content = source.read(name)
            info = zipfile.ZipInfo(filename=name, date_time=_ZIP_PINNED_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0  # MS-DOS/FAT — OS-independent
            info.external_attr = 0o600 << 16
            info.internal_attr = 0
            dest.writestr(info, content)
    source.close()
    return out.getvalue()


def _canonical_input(snapshot: ReportSnapshot, options: RendererOptions) -> bytes:
    payload = {
        "rendererId": RENDERER_ID,
        "rendererVersion": RENDERER_VERSION,
        "options": dataclasses.asdict(options),
        "report": dataclasses.asdict(snapshot.report),
        "audit": dataclasses.asdict(snapshot.audit),
        "tenant": dataclasses.asdict(snapshot.tenant),
    }
    return json.dumps(
        payload, sort_keys=True, ensure_ascii=True, separators=(",", ":"), default=_json_default
    ).encode("utf-8")


def _json_default(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.astimezone(timezone.utc).isoformat()
    if isinstance(value, (set, frozenset, tuple)):
        return list(value)
    return str(value)


def _semantic_digest(snapshot: ReportSnapshot, options: RendererOptions, container: bytes) -> str:
    """A container-independent digest over the canonical input plus the rendered
    body parts (``document.xml`` + ``styles.xml``). Proves logical identity even
    if the ZIP bytes were to vary across a future library/zlib build."""
    hasher = hashlib.sha256()
    hasher.update(_canonical_input(snapshot, options))
    with zipfile.ZipFile(io.BytesIO(container), "r") as archive:
        for part in ("word/document.xml", "word/styles.xml"):
            try:
                hasher.update(b"\x00")
                hasher.update(part.encode("ascii"))
                hasher.update(archive.read(part))
            except KeyError:
                continue
    return hasher.hexdigest()


__all__ = ["DocxRenderer", "RENDERER_ID", "RENDERER_VERSION", "CONTENT_TYPE"]
